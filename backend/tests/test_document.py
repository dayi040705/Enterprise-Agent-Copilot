"""测试: 多格式文档解析"""
import tempfile
import os


def test_read_txt():
    """纯文本文件解析"""
    from services.document import read_txt

    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
        f.write("这是测试内容。\n第二行。")
        tmp = f.name

    try:
        pages = read_txt(tmp)
        assert len(pages) == 1
        assert "这是测试内容" in pages[0]["text"]
        assert pages[0]["page"] == 1
    finally:
        os.unlink(tmp)


def test_read_docx():
    """Word 文档解析"""
    from services.document import read_docx
    from docx import Document

    doc = Document()
    doc.add_paragraph("第一条 测试制度内容。")
    doc.add_paragraph("第二条 适用于全体员工。")
    tmp = tempfile.mktemp(suffix=".docx")
    doc.save(tmp)

    try:
        pages = read_docx(tmp)
        assert len(pages) >= 1
        combined = "".join(p["text"] for p in pages)
        assert "第一条" in combined
        assert "第二条" in combined
    finally:
        os.unlink(tmp)


def test_read_md():
    """Markdown 文件解析"""
    from services.document import read_md

    with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False, encoding="utf-8") as f:
        f.write("# 标题\n\n这是 Markdown 内容。\n\n- 列表项 1\n- 列表项 2")
        tmp = f.name

    try:
        pages = read_md(tmp)
        assert len(pages) == 1
        assert "Markdown" in pages[0]["text"]
    finally:
        os.unlink(tmp)


def test_load_document_dispatcher():
    """自动格式识别"""
    from services.document import load_document

    # txt
    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False, encoding="utf-8") as f:
        f.write("测试")
        tmp_txt = f.name

    try:
        pages = load_document(tmp_txt)
        assert len(pages) == 1
    finally:
        os.unlink(tmp_txt)

    # 不支持格式应抛异常
    try:
        load_document("test.xyz")
        assert False, "应该抛异常"
    except Exception:
        pass
